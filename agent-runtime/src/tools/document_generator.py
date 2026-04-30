"""
Document Generator - 文档生成工具
基于 python-docx 生成 Word 文档，支持表格、标题、列表等格式

参考：James超级助理的 docx.md skill（使用 docx-js npm 包）
本工具提供 Python 端 Word 文档生成能力
"""
import os
import json
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 输出目录
DOCUMENTS_OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "generated_docs"
)


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table):
    """为表格设置边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


class DocumentGenerator:
    """
    Word 文档生成器
    
    支持：
    - 多级标题（Heading 1/2/3）
    - 正文段落
    - 表格生成（含表头着色）
    - 有序/无序列表
    - 图片插入
    """
    
    def __init__(self, title: str = "文档"):
        self.doc = Document()
        self.title = title
        
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加标题
        self.add_title(title)
    
    def add_title(self, text: str, level: int = 0):
        """添加文档标题"""
        if level == 0:
            heading = self.doc.add_heading(text, level=0)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        else:
            self.doc.add_heading(text, level=level)
    
    def add_heading(self, text: str, level: int = 1):
        """添加章节标题"""
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text: str, bold: bool = False, 
                      alignment: str = "left") -> None:
        """添加段落"""
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        para = self.doc.add_paragraph()
        para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
    
    def add_table(self, headers: List[str], rows: List[List[str]],
                  title: str = None, header_color: str = "1A56DB") -> None:
        """
        添加表格
        
        Args:
            headers: 表头列表
            rows: 数据行列表（每行为字符串列表）
            title: 表格标题（可选）
            header_color: 表头背景色（十六进制）
        """
        if title:
            para = self.doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # 创建表格
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shading(cell, header_color)
        
        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                cell = cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.size = Pt(10)
                # 交替行背景色
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, "F2F6FC")
        
        # 设置表格边框
        _set_table_borders(table)
        
        # 空行
        self.doc.add_paragraph()
    
    def add_bullet_list(self, items: List[str], bold_prefix: bool = False):
        """添加无序列表"""
        for item in items:
            p = self.doc.add_paragraph(style='List Bullet')
            if bold_prefix and '：' in item:
                prefix, rest = item.split('：', 1)
                run = p.add_run(prefix + '：')
                run.bold = True
                run.font.size = Pt(11)
                run = p.add_run(rest)
                run.font.size = Pt(11)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    
    def add_numbered_list(self, items: List[str]):
        """添加有序列表"""
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph(style='List Number')
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    def add_section_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_metadata(self, author: str = "", subject: str = ""):
        """添加文档元数据"""
        self.doc.core_properties.author = author or "Skill Platform"
        self.doc.core_properties.subject = subject or self.title
        self.doc.core_properties.created = datetime.now()
    
    def save(self, filename: str = None, output_dir: str = None) -> str:
        """
        保存文档到文件
        
        Args:
            filename: 文件名（不含路径），默认根据标题生成
            output_dir: 输出目录，默认使用内置目录
            
        Returns:
            保存的文件路径
        """
        output_dir = output_dir or DOCUMENTS_OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)
        
        if not filename:
            safe_title = self.title.replace(' ', '_').replace('/', '_')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.docx"
        
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = os.path.join(output_dir, filename)
        self.doc.save(filepath)
        return filepath


def generate_process_document(
    process_name: str,
    process_data: Dict[str, Any],
    output_dir: str = None,
) -> str:
    """
    生成流程分析 Word 文档
    
    Args:
        process_name: 流程名称
        process_data: 流程数据（节点、表格等）
        output_dir: 输出目录
        
    Returns:
        生成的文档路径
    """
    gen = DocumentGenerator(title=process_name)
    gen.add_metadata(author="Process Analysis Skill", subject=process_name)
    
    # 流程概述
    gen.add_heading("流程概述", level=1)
    if "description" in process_data:
        gen.add_paragraph(process_data["description"])
    
    # 流程节点表格
    nodes = process_data.get("nodes", [])
    if nodes:
        gen.add_heading("流程节点明细", level=1)
        
        # 提取所有可能的字段作为表头
        headers = ["节点名称", "类型", "负责角色", "输入", "输出", "风险点"]
        rows = []
        for node in nodes:
            rows.append([
                node.get("name", ""),
                node.get("type", ""),
                node.get("responsible_role", ""),
                ", ".join(node.get("inputs", [])) if isinstance(node.get("inputs"), list) else str(node.get("inputs", "")),
                ", ".join(node.get("outputs", [])) if isinstance(node.get("outputs"), list) else str(node.get("outputs", "")),
                ", ".join(node.get("risks", [])) if isinstance(node.get("risks"), list) else str(node.get("risks", "")),
            ])
        
        gen.add_table(
            headers=headers,
            rows=rows,
            title="流程节点列表",
            header_color="2563EB",
        )
    
    # 优化建议
    suggestions = process_data.get("optimization_suggestions", [])
    if suggestions:
        gen.add_heading("优化建议", level=1)
        for sug in suggestions:
            gen.add_paragraph(
                f"目标节点: {sug.get('target_node', '')} | 优先级: {sug.get('priority', '')}",
                bold=True,
            )
            gen.add_paragraph(sug.get("suggestion", ""))
    
    return gen.save(filename=f"{process_name}.docx", output_dir=output_dir)


def generate_report_document(
    title: str,
    sections: List[Dict[str, Any]],
    output_dir: str = None,
) -> str:
    """
    生成报告类 Word 文档
    
    Args:
        title: 文档标题
        sections: 章节列表，每个章节格式：
            {
                "heading": "章节标题",
                "level": 1,          # 标题层级
                "paragraphs": ["段落文本..."],
                "table": {
                    "title": "表格标题",
                    "headers": ["列1", "列2"],
                    "rows": [["数据1", "数据2"]]
                },
                "bullet_list": ["列表项1", "列表项2"],
                "numbered_list": ["列表项1", "列表项2"],
            }
        output_dir: 输出目录
        
    Returns:
        生成的文档路径
    """
    gen = DocumentGenerator(title=title)
    gen.add_metadata(author="Document Generation Skill", subject=title)
    
    for section in sections:
        # 添加标题
        if "heading" in section:
            gen.add_heading(section["heading"], level=section.get("level", 1))
        
        # 添加段落
        for para_text in section.get("paragraphs", []):
            gen.add_paragraph(para_text)
        
        # 添加表格
        if "table" in section:
            tbl = section["table"]
            gen.add_table(
                headers=tbl.get("headers", []),
                rows=tbl.get("rows", []),
                title=tbl.get("title"),
            )
        
        # 添加无序列表
        if "bullet_list" in section:
            gen.add_bullet_list(section["bullet_list"])
        
        # 添加有序列表
        if "numbered_list" in section:
            gen.add_numbered_list(section["numbered_list"])
    
    return gen.save(filename=f"{title}.docx", output_dir=output_dir)


def get_download_url(filename: str) -> str:
    """获取文件的下载 URL"""
    return f"/tools/download/{filename}"


def list_generated_docs(output_dir: str = None) -> List[Dict[str, str]]:
    """列出已生成的文档"""
    output_dir = output_dir or DOCUMENTS_OUTPUT_DIR
    if not os.path.isdir(output_dir):
        return []
    
    docs = []
    for fname in os.listdir(output_dir):
        if fname.endswith(('.docx', '.xlsx')):
            filepath = os.path.join(output_dir, fname)
            stat = os.stat(filepath)
            docs.append({
                "filename": fname,
                "path": filepath,
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "download_url": get_download_url(fname),
            })
    
    # 按修改时间倒序
    docs.sort(key=lambda d: d["created"], reverse=True)
    return docs
